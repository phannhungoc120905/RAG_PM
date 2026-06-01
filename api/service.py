from __future__ import annotations

import json
import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any
from unittest import result

import httpx
import time
from fastapi import HTTPException, status
from sqlalchemy import Select, func, or_, select
from sqlalchemy.orm import Session, selectinload

from config import settings
from db.models import ChunkMetadata, Document, SummaryHistory, SystemLog, User
from logger import get_logger
from ocr.runtime import ocr_service


log = get_logger("api.service")


def _clean_mindmap_node(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    if not cleaned:
        return ""
    cleaned = cleaned.replace("[", "(").replace("]", ")")
    if len(cleaned) > 80:
        cleaned = cleaned[:77].rstrip() + "..."
    return cleaned


def _build_mermaid_mindmap(root_label: str, nodes: list[str]) -> str:
    """
    Tạo Mermaid mindmap với multi-level hierarchy (thay vì flat list).
    Chia nodes thành 3 nhóm chính để sơ đồ không quá xấp xủa.
    """
    lines = ["mindmap", f"  root(({root_label}))"]
    
    # Chia nodes thành 3 category (thường dùng cho tài liệu hành chính):
    # - Thông tin chung / Mục đích
    # - Nội dung chính / Điều khoản
    # - Thực hiện / Lưu ý
    
    if len(nodes) <= 5:
        # Nếu nodes ít, display flat
        for node in nodes:
            lines.append(f"    {node}")
    else:
        # Chia thành 3 group có hierarchy
        group_size = (len(nodes) + 2) // 3
        
        # Group 1: Thông tin cơ bản
        lines.append("    Thông tin")
        for node in nodes[:group_size]:
            lines.append(f"      {node}")
        
        # Group 2: Nội dung chính
        if len(nodes) > group_size:
            lines.append("    Nội dung")
            for node in nodes[group_size:2*group_size]:
                lines.append(f"      {node}")
        
        # Group 3: Thực hiện / Lưu ý
        if len(nodes) > 2*group_size:
            lines.append("    Thực hiện")
            for node in nodes[2*group_size:]:
                lines.append(f"      {node}")
    
    return "\n".join(lines)


def get_supported_formats() -> list[str]:
    return ["pdf", "docx", "txt", "jpg", "jpeg", "png", "bmp", "tif", "tiff"]


async def check_ollama_health() -> dict[str, Any]:

    try:
        async with httpx.AsyncClient() as client:
            response = await client.get(settings.OLLAMA_TAGS_URL, timeout=10.0)
            response.raise_for_status()
            payload = response.json()
        models = [item.get("name") for item in payload.get("models", []) if item.get("name")]
        return {
            "ok": True,
            "base_url": settings.OLLAMA_BASE_URL,
            "model_name": settings.MODEL_NAME,
            "available_models": models,
        }
    except Exception as exc:
        return {
            "ok": False,
            "base_url": settings.OLLAMA_BASE_URL,
            "model_name": settings.MODEL_NAME,
            "error": str(exc),
        }


async def generate_mindmap_for_document(
    db: Session,
    *,
    current_user: User,
    document_id: int,
    use_llm: bool = True,
    max_nodes: int = 12,
) -> dict[str, Any]:
    document = get_document_detail(db, document_id=document_id, current_user=current_user)
    root_label = document.document_title or document.original_filename or f"Van ban #{document.id}"

    if use_llm:
        query = (
            "=== HƯỚNG DẪN TẠO SƠ ĐỒ TƯ DUY (MINDMAP) ===\n\n"
            "Nhiệm vụ: Phân tích tài liệu và tạo sơ đồ tư duy MERMAID.JS với CẤU TRÚC PHÂN CẤP rõ ràng.\n\n"
            "ĐỊNH DẠNG YÊU CẦU (bắt đầu bằng 'mindmap'):\n"
            "mindmap\n"
            "  root((Tiêu đề tài liệu))\n"
            "    Nhóm 1: Thông tin cơ bản\n"
            "      Chi tiết 1\n"
            "      Chi tiết 2\n"
            "    Nhóm 2: Nội dung chính\n"
            "      Điểm 1\n"
            "      Điểm 2\n"
            "    Nhóm 3: Thực hiện / Kết luận\n"
            "      Bước 1\n"
            "      Bước 2\n\n"
            "YÊUQCẦU:\n"
            "1. Tạo tối đa 3 nhóm (level 1), mỗi nhóm tối đa 4 items (level 2)\n"
            "2. Không thêm level 3 trở lên\n"
            "3. Văn bản mỗi node ngắn gọn (max 50 ký tự)\n"
            "4. Chỉ trả về mã MERMAID, không giải thích thêm"
        )
        result = await ocr_service.get_rag_answer(
            query,
            top_k=8,
            document_ids=[document.id],
        )
        answer = (result.get("answer") or "").replace("```mermaid", "").replace("```", "").strip()
        if answer.startswith("mindmap"):
            return {"mode": "llm", "mermaid": answer, "source": "llm"}

    chunks = db.scalars(
        select(ChunkMetadata)
        .where(ChunkMetadata.document_id == document.id)
        .order_by(
            (ChunkMetadata.page_number.is_(None)).asc(),
            ChunkMetadata.page_number.asc(),
            (ChunkMetadata.chunk_index.is_(None)).asc(),
            ChunkMetadata.chunk_index.asc(),
        )
    ).all()

    nodes: list[str] = []
    seen: set[str] = set()
    for chunk in chunks:
        candidate = chunk.section_title or chunk.section_code or chunk.content_preview or ""
        cleaned = _clean_mindmap_node(candidate)
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        nodes.append(cleaned)
        if len(nodes) >= max_nodes:
            break

    if not nodes:
        nodes = ["Khong the trich xuat cau truc", "Vui long thu lai"]

    mermaid = _build_mermaid_mindmap(root_label, nodes)
    return {"mode": "fallback", "mermaid": mermaid, "source": "chunks"}


def list_documents(
    db: Session,
    *,
    current_user: User,
    page: int,
    page_size: int,
    search: str | None = None,
    document_type: str | None = None,
    processing_status: str | None = None,
    review_status: str | None = None,
) -> tuple[list[Document], int]:
    statement: Select[tuple[Document]] = select(Document).options(
        selectinload(Document.owner),
        selectinload(Document.summaries),
    ).where(Document.deleted_at.is_(None))
    count_statement = select(func.count()).select_from(Document).where(Document.deleted_at.is_(None))

    if current_user.role != "admin" and _get_permission_group_code(current_user) != "AGENCY_LEADER":
        statement = statement.where(Document.owner_id == current_user.id)
        count_statement = count_statement.where(Document.owner_id == current_user.id)

    if search:
        criteria = or_(
            Document.original_filename.ilike(f"%{search}%"),
            Document.document_title.ilike(f"%{search}%"),
            Document.document_number.ilike(f"%{search}%"),
            Document.document_type.ilike(f"%{search}%"),
            Document.classification_label.ilike(f"%{search}%"),
        )
        statement = statement.where(criteria)
        count_statement = count_statement.where(criteria)

    if document_type:
        statement = statement.where(Document.document_type == document_type)
        count_statement = count_statement.where(Document.document_type == document_type)
    if processing_status:
        statement = statement.where(Document.processing_status == processing_status)
        count_statement = count_statement.where(Document.processing_status == processing_status)
    if review_status:
        statement = statement.where(Document.review_status == review_status)
        count_statement = count_statement.where(Document.review_status == review_status)

    total = db.scalar(count_statement) or 0
    items = db.scalars(
        statement.order_by(Document.id.desc()).offset((page - 1) * page_size).limit(page_size)
    ).all()
    return items, total


def get_document_detail(db: Session, document_id: int, current_user: User) -> Document:
    document = db.scalar(
        select(Document)
        .options(
            selectinload(Document.owner),
            selectinload(Document.chunks),
            selectinload(Document.summaries),
        )
        .where(Document.id == document_id, Document.deleted_at.is_(None))
    )
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    _assert_document_access(document, current_user)
    return document


async def upload_document(
    db: Session,
    *,
    current_user: User,
    file_bytes: bytes,
    original_filename: str,
    mime_type: str | None = None,
    auto_summary: bool = False,
) -> tuple[Document, SummaryHistory | None]:
    final_path: Path | None = None     
    embeddings_written = False        
    if not original_filename or "." not in original_filename:
        raise HTTPException(status_code=400, detail="Invalid filename")    
    extension = original_filename.rsplit(".", 1)[-1].lower()
    if extension not in get_supported_formats():
        raise HTTPException(status_code=400, detail=f"Unsupported file format: {extension}")

    processing_path = _write_upload_file(file_bytes, original_filename, "processing")
    document = Document(
        filename=processing_path.name,
        original_filename=original_filename,
        file_path=str(processing_path),
        file_size_kb=max(1, len(file_bytes) // 1024) if file_bytes else 0,
        status="processing",
        source_format=extension,
        mime_type=mime_type,
        owner_id=current_user.id,
        uploaded_by=current_user.id,
        processing_status="processing",
        review_status="pending",
        created_at=datetime.now(),
    )
    db.add(document)
    db.commit()
    db.refresh(document)

    try:
        started_at = time.perf_counter()
        result = ocr_service.process_document(file_bytes, original_filename)
        ocr_elapsed = time.perf_counter() - started_at
        log.info(
            "upload_ocr_completed",
            extra={"document_id": document.id, "filename": original_filename, "seconds": round(ocr_elapsed, 3)},
        )
        file_size_mb = len(file_bytes) / (1024 * 1024) if file_bytes else 0
        image_exts = {"jpg", "jpeg", "png", "bmp", "tif", "tiff"}
        is_image = (extension in image_exts)
        should_fix_with_llm = (
            settings.OCR_FIX_WITH_LLM and not is_image and file_size_mb <= settings.OCR_FIX_WITH_LLM_MAX_FILE_MB
        )
        if should_fix_with_llm:
            fix_started_at = time.perf_counter()
            result = await ocr_service.fix_processed_result_with_llm(result)
            log.info(
                "upload_ocr_fixed_with_llm",
                extra={
                    "document_id": document.id,
                    "filename": original_filename,
                    "seconds": round(time.perf_counter() - fix_started_at, 3),
                    "file_size_mb": round(file_size_mb, 2),
                },
            )
        else:
            log.info(
                "upload_ocr_skip_llm",
                extra={
                    "document_id": document.id,
                    "filename": original_filename,
                    "file_size_mb": round(file_size_mb, 2),
                    "threshold_mb": settings.OCR_FIX_WITH_LLM_MAX_FILE_MB,
                    "enabled": settings.OCR_FIX_WITH_LLM,
                },
            )
        print("DEBUG chunks type:", type(result.get("chunks", [])))
        print("DEBUG chunks sample:", str(result.get("chunks", []))[:300])

        embed_started_at = time.perf_counter()
        final_path = _move_upload_file(processing_path, "done")
        enriched_chunks = _build_enriched_chunks(document.id, result.get("chunks", []))
        start_vector_id = int(ocr_service.index.ntotal)
        embedded_chunks = ocr_service.embed_chunks(enriched_chunks)
        ocr_service.store_embeddings(embedded_chunks)
        log.info(
            "upload_embedding_completed",
            extra={"document_id": document.id, "filename": original_filename, "seconds": round(time.perf_counter() - embed_started_at, 3)},
        )
        embeddings_written = True

        for index, chunk in enumerate(enriched_chunks):
            metadata = chunk.get("metadata", {})
            db.add(
                ChunkMetadata(
                    document_id=document.id,
                    chunk_index=index,
                    chunk_type=metadata.get("section_type") or "semantic",
                    section_type=metadata.get("section_type"),
                    section_code=_section_code(metadata),
                    section_title=metadata.get("anchor_text"),
                    page_number=metadata.get("page_number"),
                    start_line=metadata.get("start_line"),
                    end_line=metadata.get("end_line"),
                    end_page=metadata.get("page_number"),
                    token_count=len(chunk.get("content", "").split()),
                    faiss_index_id=start_vector_id + index,
                    embedding_status="completed",
                    embedding_model=settings.EMBEDDING_MODEL_NAME,
                    bm25_text=chunk.get("content"),
                    citation_json=_json_dumps(
                        {
                            "page_number": metadata.get("page_number"),
                            "page_label": metadata.get("page_label"),
                            "start_line": metadata.get("start_line"),
                            "end_line": metadata.get("end_line"),
                            "dieu": metadata.get("dieu"),
                            "khoan": metadata.get("khoan"),
                        }
                    ),
                    metadata_json=_json_dumps(metadata),
                    content_preview=chunk.get("content", "")[:400] or None,
                )
            )

        document.filename = final_path.name
        document.file_path = str(final_path)
        document.status = "processed"
        document.document_type = result["classification"].get("document_type")
        document.document_number = result["structure"].get("document_code")
        document.document_title = result["structure"].get("summary") or original_filename
        document.page_count = result.get("page_count")
        document.ocr_text = result.get("raw_text")
        document.clean_text = result.get("clean_text")
        document.processing_status = "completed"
        document.processing_error = None
        document.classification_label = result["classification"].get("document_type")
        document.classification_score = result["classification"].get("confidence")
        document.structure_json = _json_dumps(result.get("structure"))
        document.page_index_json = _json_dumps(result.get("page_index"))
        document.storage_meta_json = _json_dumps({"supported_formats": result.get("supported_formats", get_supported_formats())})
        document.processed_by = current_user.id
        document.processed_at = datetime.now()
        document.updated_at = datetime.now()
        db.add(document)
        db.commit()
        db.refresh(document)
        log.info(
            "upload_document_completed",
            extra={
                "document_id": document.id,
                "filename": original_filename,
                "total_seconds": round(time.perf_counter() - started_at, 3),
            },
        )

        _write_system_log(
            db,
            user_id=current_user.id,
            action="document_upload",
            detail=_json_dumps({"document_id": document.id, "filename": original_filename}),
            module_name="api.documents",
            entity_type="document",
            entity_id=document.id,
            log_type="usage",
        )
    except Exception as e:
        db.rollback()
        if embeddings_written:
            try:
                _rebuild_vector_storage(db)
            except Exception:
                log.exception("upload_vector_rebuild_failed", extra={"document_id": document.id})
        if final_path and final_path.exists():
            try:
                failed_path = _move_upload_file(final_path, "failed")
                document.file_path = str(failed_path)
                document.filename = failed_path.name
            except Exception:
                log.exception("upload_failed_file_move_failed", extra={"document_id": document.id, "filename": original_filename})
        log.error("upload_processing_failed", extra={"document_id": document.id, "error": str(e)})
        document.processing_status = "failed"
        document.status = "failed"
        document.processing_error = str(e)
        document.updated_at = datetime.now()
        db.add(document)
        db.commit()
        log.exception(
            "document_upload_failed",
          extra={
            "document_id": document.id,
            "filename": original_filename,
            "error": str(e),
          }
        )
        raise HTTPException(status_code=500, detail=f"Lỗi xử lý tài liệu: {str(e)}")
    summary_record = None
    if auto_summary:
        try:
            summary_record = await create_summary_for_document(
                db,
                document_id=document.id,
                current_user=current_user,
            )
        except Exception:
            log.exception(
                "document_auto_summary_failed",
                extra ={
                    "document_id": document.id,
                    "filename": original_filename,
                }
            )

    return document, summary_record


async def create_summary_for_document(
    db: Session,
    *,
    document_id: int,
    current_user: User,
    title: str | None = None,
    prompt_template: str = "default_admin_summary",
) -> SummaryHistory:
    document = get_document_detail(db, document_id, current_user)
    if not document.clean_text:
        raise HTTPException(status_code=400, detail="Document has no extracted text to summarize")

    chunk_ids = [chunk.id for chunk in sorted(document.chunks, key=lambda item: item.chunk_index or 0)]
    prompt = _build_summary_prompt(document)
    summary_text = await _generate_llm_text(prompt)
    if not summary_text:
        raise HTTPException(status_code=502, detail="Model did not return a summary")
    
    # Enforce quality on raw LLM output: remove prompt leakage, headings, hallucinations.
    summary_text = _postprocess_summary(summary_text)
    if _looks_like_refusal(summary_text):
        summary_text = _fallback_summary_from_source(document.clean_text or "")
    elif _looks_like_source_copy(summary_text, document.clean_text or ""):
        summary_text = _fallback_summary_from_source(document.clean_text or "")

    version_no = (db.scalar(select(func.max(SummaryHistory.version_no)).where(SummaryHistory.document_id == document.id)) or 0) + 1
    summary = SummaryHistory(
        document_id=document.id,
        user_id=current_user.id,
        summary_type="summary",
        version_no=version_no,
        title=title or document.document_title or document.original_filename,
        summary_text=summary_text,
        prompt_template=prompt_template,
        model_name=settings.MODEL_NAME,
        source_chunk_ids_json=_json_dumps(chunk_ids),
        groundedness_score=1.0 if chunk_ids else 0.0,
        hallucination_flag=False,
        is_reviewed=False,
        created_at=datetime.now(),
        updated_at=datetime.now(),
    )
    document.document_summary = summary_text
    document.updated_at = datetime.now()
    db.add(summary)
    db.add(document)
    db.commit()
    db.refresh(summary)

    _write_system_log(
        db,
        user_id=current_user.id,
        action="document_summarize",
        detail=_json_dumps({"document_id": document.id, "summary_id": summary.id, "version_no": version_no}),
        module_name="api.summaries",
        entity_type="summary",
        entity_id=summary.id,
        log_type="usage",
    )
    return summary


async def answer_question(
    db: Session,
    *,
    current_user: User,
    query: str,
    top_k: int = 5,
    document_ids: list[int] | None = None,
) -> dict[str, Any]:
    allowed_document_ids = _resolve_allowed_document_ids(db, current_user, document_ids)
    db_answer = await _answer_question_from_db_chunks(
        db,
        query=query,
        top_k=top_k,
        document_ids=allowed_document_ids,
    )
    if db_answer is not None:
        return db_answer
    return await ocr_service.get_rag_answer(query, top_k=top_k, document_ids=allowed_document_ids)


async def _answer_question_from_db_chunks(
    db: Session,
    *,
    query: str,
    top_k: int,
    document_ids: list[int] | None,
) -> dict[str, Any] | None:
    query = (query or "").strip()
    if not query:
        return None

    statement = (
        select(ChunkMetadata)
        .join(Document, ChunkMetadata.document_id == Document.id)
        .where(Document.deleted_at.is_(None))
        .order_by(ChunkMetadata.document_id.desc(), ChunkMetadata.chunk_index.asc())
        .limit(600)
    )
    if document_ids is not None:
        if not document_ids:
            return {
                "answer": "Khong tim thay thong tin phu hop trong tai lieu.",
                "sources": [],
                "source_chunks": [],
                "grounded": False,
            }
        statement = statement.where(ChunkMetadata.document_id.in_(document_ids))

    rows = list(db.scalars(statement).all())
    if not rows:
        return None

    query_terms = _important_terms(query)
    wants_summary = _is_summary_query(query)
    scored: list[tuple[int, int, ChunkMetadata]] = []
    for row in rows:
        content = row.bm25_text or row.content_preview or ""
        if not content.strip():
            continue
        lowered = content.lower()
        content_terms = set(re.findall(r"\w+", lowered, flags=re.UNICODE))
        overlap = len(query_terms.intersection(content_terms))
        number_bonus = min(len(re.findall(r"\d+", content)), 4)
        section_bonus = 2 if row.section_title or row.section_code else 0
        summary_bonus = min(len(content.split()) // 80, 4) if wants_summary else 0
        score = overlap * 6 + number_bonus + section_bonus + summary_bonus
        if wants_summary or score > 0:
            scored.append((score, row.chunk_index or 0, row))

    if not scored:
        return {
            "answer": "Khong tim thay thong tin phu hop trong tai lieu.",
            "sources": [],
            "source_chunks": [],
            "grounded": False,
        }

    scored.sort(key=lambda item: (item[0], -item[1]), reverse=True)
    selected_rows = [row for _, _, row in scored[: min(max(top_k, 3), 6)]]
    chunks = [_chunk_row_to_rag_chunk(row) for row in selected_rows]
    prompt = ocr_service.build_grounded_prompt(query, chunks)

    try:
        answer = await _generate_llm_text(prompt)
        if not answer or _looks_like_source_copy(answer, " ".join(chunk["content"] for chunk in chunks)):
            answer = ocr_service._fallback_rag_answer_from_chunks(query, chunks)
        grounded = ocr_service.validate_answer_vs_context(answer, chunks)
    except Exception:
        answer = ocr_service._fallback_rag_answer_from_chunks(query, chunks)
        grounded = False

    return {
        "answer": answer,
        "sources": [chunk["content"] for chunk in chunks],
        "source_chunks": chunks,
        "grounded": grounded,
    }


def _chunk_row_to_rag_chunk(row: ChunkMetadata) -> dict[str, Any]:
    metadata = _json_loads(row.metadata_json, default={})
    metadata.setdefault("document_id", row.document_id)
    metadata.setdefault("chunk_index", row.chunk_index)
    metadata.setdefault("page_number", row.page_number)
    metadata.setdefault("start_line", row.start_line)
    metadata.setdefault("end_line", row.end_line)
    metadata.setdefault("citation_anchor", _build_citation_anchor(metadata))
    return {
        "content": row.bm25_text or row.content_preview or "",
        "metadata": metadata,
        "final_score": None,
    }


def _important_terms(text: str) -> set[str]:
    stopwords = {
        "cho", "toi", "biet", "hay", "neu", "thi", "la", "cua", "cac", "nhung",
        "trong", "ngoai", "the", "nao", "gi", "ve", "voi", "mot", "duoc", "khong",
        "co", "noi", "dung", "tai", "lieu", "tom", "tat",
    }
    return {
        token
        for token in re.findall(r"\w+", text.lower(), flags=re.UNICODE)
        if len(token) >= 3 and token not in stopwords
    }


def _is_summary_query(query: str) -> bool:
    lowered = query.lower()
    return any(
        signal in lowered
        for signal in ("tom tat", "tóm tắt", "noi dung chinh", "nội dung chính", "khai quat", "khái quát")
    )


def search_chunks(
    db: Session,
    *,
    current_user: User,
    query: str,
    top_k: int = 5,
    document_ids: list[int] | None = None,
) -> list[dict[str, Any]]:
    allowed_document_ids = _resolve_allowed_document_ids(db, current_user, document_ids)
    return ocr_service.hybrid_search(query, top_k=top_k, threshold=0.0, document_ids=allowed_document_ids)


def list_history(
    db: Session,
    *,
    current_user: User | None,
    page: int,
    page_size: int,
    document_id: int | None = None,
) -> tuple[list[SummaryHistory], int]:
    statement: Select[tuple[SummaryHistory]] = select(SummaryHistory).options(
        selectinload(SummaryHistory.document),
        selectinload(SummaryHistory.user),
    ).where(SummaryHistory.is_deleted.is_(False))
    count_statement = select(func.count()).select_from(SummaryHistory).where(SummaryHistory.is_deleted.is_(False))

    if (
        current_user
        and current_user.role != "admin"
        and _get_permission_group_code(current_user) != "AGENCY_LEADER"
    ):
        statement = statement.join(Document, SummaryHistory.document_id == Document.id).where(Document.owner_id == current_user.id)
        count_statement = count_statement.join(Document, SummaryHistory.document_id == Document.id).where(Document.owner_id == current_user.id)
    if document_id is not None:
        statement = statement.where(SummaryHistory.document_id == document_id)
        count_statement = count_statement.where(SummaryHistory.document_id == document_id)

    total = db.scalar(count_statement) or 0
    items = db.scalars(
        statement.order_by(SummaryHistory.id.desc()).offset((page - 1) * page_size).limit(page_size)
    ).all()
    return items, total


def get_summary_detail(db: Session, summary_id: int, current_user: User) -> SummaryHistory:
    summary = db.scalar(
        select(SummaryHistory)
        .options(
            selectinload(SummaryHistory.document).selectinload(Document.chunks),
            selectinload(SummaryHistory.user),
        )
        .where(SummaryHistory.id == summary_id, SummaryHistory.is_deleted.is_(False))
    )
    if not summary:
        raise HTTPException(status_code=404, detail="Summary not found")
    _assert_document_access(summary.document, current_user)
    return summary


def review_summary(
    db: Session,
    *,
    summary_id: int,
    current_user: User,
    approved: bool,
    note: str | None,
) -> SummaryHistory:
    if current_user.role != "admin" and _get_permission_group_code(current_user) != "AGENCY_LEADER":
        raise HTTPException(status_code=403, detail="Agency leader or admin access required for review")

    summary = get_summary_detail(db, summary_id, current_user)
    summary.is_reviewed = approved
    summary.reviewed_by = current_user.id
    summary.review_note = note
    summary.updated_at = datetime.now()
    summary.document.review_status = "approved" if approved else "needs_revision"
    summary.document.reviewed_by = current_user.id
    summary.document.reviewed_at = datetime.now()
    summary.document.updated_at = datetime.now()
    db.add(summary)
    db.add(summary.document)
    db.commit()
    db.refresh(summary)
    return summary


def update_summary_text(
    db: Session,
    *,
    summary_id: int,
    current_user: User,
    summary_text: str,
    title: str | None = None,
    note: str | None = None,
) -> SummaryHistory:
    summary = get_summary_detail(db, summary_id, current_user)
    summary.summary_text = summary_text
    if title is not None:
        summary.title = title
    if note is not None:
        summary.review_note = note
    summary.updated_at = datetime.now()
    summary.document.document_summary = summary_text
    summary.document.updated_at = datetime.now()
    db.add(summary)
    db.add(summary.document)
    db.commit()
    db.refresh(summary)
    return summary


def leave_summary_feedback(
    db: Session,
    *,
    summary_id: int,
    current_user: User,
    score: int | None,
    comment: str | None,
) -> SummaryHistory:
    summary = get_summary_detail(db, summary_id, current_user)
    summary.feedback_score = score
    summary.feedback_comment = comment
    summary.updated_at = datetime.now()
    db.add(summary)
    db.commit()
    db.refresh(summary)
    return summary


def export_summary(db: Session, *, summary_id: int, current_user: User, export_format: str) -> Path:
    summary = get_summary_detail(db, summary_id, current_user)
    export_format = export_format.lower()
    export_dir = Path(settings.BACKUP_DIR) / "exports"
    export_dir.mkdir(parents=True, exist_ok=True)
    safe_title = _slugify_filename(summary.title or f"summary_{summary.id}")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if export_format == "txt":
        export_path = export_dir / f"{safe_title}_{timestamp}.txt"
        export_path.write_text(summary.summary_text or "", encoding="utf-8")
    elif export_format == "json":
        export_path = export_dir / f"{safe_title}_{timestamp}.json"
        export_path.write_text(
            _json_dumps(
                {
                    "summary_id": summary.id,
                    "document_id": summary.document_id,
                    "title": summary.title,
                    "summary_text": summary.summary_text,
                    "model_name": summary.model_name,
                    "created_at": summary.created_at,
                }
            ),
            encoding="utf-8",
        )
    elif export_format == "docx":
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise HTTPException(status_code=501, detail="python-docx is required for DOCX export") from exc
        export_path = export_dir / f"{safe_title}_{timestamp}.docx"
        word_doc = WordDocument()
        word_doc.add_heading(summary.title or "Summary", level=1)
        word_doc.add_paragraph(summary.summary_text or "")
        word_doc.save(export_path)
    else:
        raise HTTPException(status_code=400, detail="Unsupported export format")

    exported_formats = _json_loads(summary.exported_formats_json, default=[])
    exported_formats.append({"format": export_format, "path": str(export_path), "exported_at": datetime.now().isoformat()})
    summary.exported_formats_json = _json_dumps(exported_formats)
    summary.updated_at = datetime.now()
    db.add(summary)
    db.commit()
    return export_path


def delete_document_cascade(db: Session, *, document_id: int, current_user: User) -> dict[str, Any]:
    document = get_document_detail(db, document_id, current_user)
    file_path = Path(document.file_path) if document.file_path else None
    payload = {"document_id": document.id, "filename": document.original_filename}
    document.deleted_at = datetime.now()
    document.updated_at = datetime.now()
    document.processing_status = "deleted"
    document.status = "deleted"
    db.add(document)
    db.commit()

    if file_path and file_path.exists():
        file_path.unlink(missing_ok=True)

    _rebuild_vector_storage(db)
    _write_system_log(
        db,
        user_id=current_user.id,
        action="document_delete",
        detail=_json_dumps(payload),
        module_name="api.documents",
        entity_type="document",
        entity_id=document_id,
        log_type="usage",
    )
    return payload


def get_latest_history_public(db: Session, limit: int = 20) -> list[dict[str, Any]]:
    items = db.scalars(
        select(SummaryHistory)
        .options(selectinload(SummaryHistory.document))
        .where(SummaryHistory.is_deleted.is_(False))
        .order_by(SummaryHistory.id.desc())
        .limit(limit)
    ).all()
    return [
        {
            "id": item.id,
            "document_id": item.document_id,
            "filename": item.document.original_filename if item.document else None,
            "title": item.title,
            "summary": item.summary_text,
            "timestamp": item.created_at.strftime("%Y-%m-%d %H:%M:%S") if item.created_at else None,
        }
        for item in items
    ]


def _resolve_allowed_document_ids(db: Session, current_user: User, document_ids: list[int] | None) -> list[int] | None:
    if current_user.role == "admin" or _get_permission_group_code(current_user) == "AGENCY_LEADER":
        return document_ids

    requested_ids = set(document_ids or [])
    owned_ids = set(
        db.scalars(select(Document.id).where(Document.owner_id == current_user.id, Document.deleted_at.is_(None))).all()
    )
    if requested_ids:
        forbidden = requested_ids - owned_ids
        if forbidden:
            raise HTTPException(status_code=403, detail="You do not have access to one or more requested documents")
        return list(requested_ids)
    return list(owned_ids)


def _assert_document_access(document: Document, current_user: User) -> None:
    if current_user.role == "admin" or _get_permission_group_code(current_user) == "AGENCY_LEADER":
        return
    if document.owner_id != current_user.id:
        raise HTTPException(status_code=403, detail="You do not have access to this document")


def _get_permission_group_code(current_user: User) -> str | None:
    if current_user.permission_group:
        return current_user.permission_group.code
    return None


def _build_enriched_chunks(document_id: int, chunks: list) -> list[dict[str, Any]]:
    enriched_chunks: list[dict[str, Any]] = []
    for index, raw_chunk in enumerate(chunks):
        if isinstance(raw_chunk, str):
            chunk = {"content": raw_chunk, "metadata": {}}
        elif isinstance(raw_chunk, dict):
            chunk = raw_chunk
        else:
            chunk = {"content": str(raw_chunk), "metadata": {}}

        metadata = dict(chunk.get("metadata", {}))
        metadata["document_id"] = document_id
        metadata["chunk_index"] = index
        metadata["citation_anchor"] = _build_citation_anchor(metadata)
        enriched_chunks.append({"content": chunk.get("content", ""), "metadata": metadata})

    return enriched_chunks

def _build_citation_anchor(metadata: dict[str, Any]) -> str:
    page = metadata.get("page_number")
    start_line = metadata.get("start_line")
    end_line = metadata.get("end_line")
    if page and start_line and end_line:
        return f"Trang {page}, dong {start_line}-{end_line}"
    if page:
        return f"Trang {page}"
    return ""


def _section_code(metadata: dict[str, Any]) -> str | None:
    dieu = metadata.get("dieu")
    khoan = metadata.get("khoan")
    if dieu and khoan:
        return f"Dieu {dieu}/Khoan {khoan}"
    if dieu:
        return f"Dieu {dieu}"
    return None


def _write_upload_file(file_bytes: bytes, original_filename: str, stage: str) -> Path:
    target_dir = Path(settings.UPLOAD_DIR) / stage
    target_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    stored_name = f"{timestamp}_{_slugify_filename(original_filename)}"
    stored_path = target_dir / stored_name
    stored_path.write_bytes(file_bytes)
    return stored_path


def _move_upload_file(path: Path, stage: str) -> Path:
    if not path.exists():
        return path
    target_dir = Path(settings.UPLOAD_DIR) / stage
    target_dir.mkdir(parents=True, exist_ok=True)
    target_path = target_dir / path.name

    # Windows deployments in this workspace can create files that are readable
    # but not deletable from the processing folder. Prefer copy + best-effort
    # cleanup so upload/summarize does not fail on stage transitions.
    if target_path.exists():
        target_path.unlink(missing_ok=True)
    shutil.copy2(path, target_path)
    try:
        path.unlink(missing_ok=True)
    except PermissionError:
        log.warning(
            "upload_stage_cleanup_skipped",
            extra={"source_path": str(path), "target_path": str(target_path), "stage": stage},
        )
    return target_path


def _slugify_filename(filename: str) -> str:
    sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", filename).strip("._")
    return sanitized or "document"


def _json_dumps(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, default=str)


def _json_loads(value: str | None, default: Any) -> Any:
    if not value:
        return default
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        return default


def _write_system_log(
    db: Session,
    *,
    action: str,
    user_id: int | None = None,
    detail: str | None = None,
    module_name: str | None = None,
    entity_type: str | None = None,
    entity_id: int | None = None,
    log_type: str = "system",
) -> None:
    db.add(
        SystemLog(
            user_id=user_id,
            action=action,
            detail=detail,
            status_code=200,
            module_name=module_name,
            entity_type=entity_type,
            entity_id=entity_id,
            log_type=log_type,
            created_at=datetime.now(),
        )
    )
    db.commit()


async def _generate_llm_text(prompt: str) -> str:
    async with httpx.AsyncClient() as client:
        response = await client.post(
            settings.OLLAMA_GENERATE_URL,
            json={
                "model": settings.MODEL_NAME,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "num_ctx": settings.OLLAMA_NUM_CTX,
                    "num_predict": 260,
                    "temperature": 0.2,
                    "repeat_penalty": 1.15,
                },
            },
            timeout=httpx.Timeout(settings.OLLAMA_TIMEOUT_SECONDS, connect=10.0),
        )
        response.raise_for_status()
        return response.json().get("response", "").strip()


def _build_summary_prompt(document: Document) -> str:
    source_text = _prepare_summary_source_text(document.clean_text or "")

    return (
        "NHIEM VU: TOM TAT NOI DUNG CHINH CUA TAI LIEU HANH CHINH.\n"
        "Khong chep lai nguyen van. Khong viet lai tung muc. Khong copy nua dau tai lieu.\n"
        "Hay doc toan bo tai lieu, rut ra y chinh, chu the, muc dich, noi dung quy dinh/yeu cau, thoi han/so lieu quan trong neu co.\n\n"
        "QUY TAC:\n"
        "- Chi dung thong tin co trong tai lieu.\n"
        "- Bo qua tieu ngu, header/footer, noi lap lai, mau bieu thuc hanh chinh khong quan trong.\n"
        "- Neu tai lieu dai, uu tien ket luan, quyet dinh, nhiem vu, doi tuong ap dung va moc thoi gian.\n"
        "- Van ban dau ra phai ngan hon tai lieu nguon rat nhieu.\n"
        "- Tuyet doi khong trich/copy lien tiep qua 25 tu tu tai lieu nguon.\n\n"
        "DINH DANG DAU RA:\n"
        "Viet 1 doan van duy nhat, 80-130 tu, 4-6 cau, khong markdown, khong bullet, khong tieu de.\n\n"
        f"TAI LIEU NGUON:\n{source_text}\n\n"
        "TOM TAT NOI DUNG CHINH:"
    )


def _prepare_summary_source_text(text: str) -> str:
    if not text:
        return ""

    # Preflight cleanup: normalize Unicode and reuse OCR normalization pipeline
    # to reduce mojibake / OCR artifacts before sending text to the LLM.
    normalized = unicodedata.normalize("NFC", text)
    try:
        normalized = ocr_service.normalize_text(normalized)
    except Exception:
        # If OCR service cleanup is unavailable, continue with local cleanup.
        pass

    normalized = _repair_common_ocr_typos(normalized)

    lines: list[str] = []
    for raw_line in normalized.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if re.fullmatch(r"(?:---\s*Page\s*\d+\s*---|Trang\s*\d+(?:\s*[-:–—].*)?)", line, flags=re.IGNORECASE):
            continue
        if re.fullmatch(r"\[[^\]]+\]", line):
            continue
        if re.fullmatch(r"[A-ZÀ-Ỵ0-9\s.,:;()/%+-]{6,}", line) and len(line) <= 80:
            # Drop short all-caps headings that commonly come from page banners.
            continue
        lines.append(line)
    collapsed = re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()
    return collapsed[:12000]


def _repair_common_ocr_typos(text: str) -> str:
    replacements = {
        "TỦNH HÌNH": "TÌNH HÌNH",
        "TỦNH HINH": "TÌNH HÌNH",
        "ban Tom Tat Dieu Hạnh": "ban tom tat dieu hanh",
        "Tom Tat Dieu Hạnh": "Tom Tat Dieu Hanh",
    }
    fixed = text
    for wrong, right in replacements.items():
        fixed = fixed.replace(wrong, right)
    return fixed


def _postprocess_summary(text: str) -> str:
    """
    Enforce strict quality on LLM output:
    1. Remove prompt leakage (lines that match system instructions).
    2. Remove hallucinated headings/recommendations not in source.
    3. Collapse into single paragraph (no line breaks between sentences).
    4. Keep only substantive content from source document.
    """
    if not text:
        log.warning("Summary postprocess: Empty input text")
        return ""

    # Strip leading/trailing whitespace
    text = text.strip()
    log.debug(f"Summary postprocess: Raw LLM output ({len(text)} chars):\n{text[:500]}")

    # Remove lines that are prompt echoes or system instruction leakage
    prompt_leakage_patterns = [
        r"^\s*Ban\s+la\s+chuyen\s+gia",  # "Ban la chuyen gia..."
        r"^\s*Toi\s+xin\s+bao\s+cao",  # "Toi xin bao cao..."
        r"^\s*Nhiem\s+vu\s+cua\s+ban",  # "Nhiem vu cua ban..."
        r"^\s*Quy\s+tac",  # "Quy tac..."
        r"^\s*Dau\s+ra\s+mong\s+muon",  # "Dau ra mong muon..."
    ]

    lines = []
    filtered_lines = {"leakage": [], "heading": [], "hallucination": []}
    
    for raw_line in text.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        # Check if line matches any prompt leakage pattern
        is_leakage = any(re.search(pat, line, re.IGNORECASE) for pat in prompt_leakage_patterns)
        if is_leakage:
            filtered_lines["leakage"].append(line[:80])
            continue

        # Remove ONLY pure all-caps short headings (very strict: must be ONLY caps/numbers/spaces, 4-50 chars)
        if re.fullmatch(r"[A-ZÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴ0-9\s—–-]{4,50}", line) and len(line) <= 50:
            filtered_lines["heading"].append(line)
            continue

        # Remove ONLY very specific hallucinated recommendation starts (not generic content)
        hallucination_patterns = [
            r"^\s*(Lời khuyên|De xuat|Kien nghi|Nhan xet|Khuyên tăng|Khuyên giảm|Nên tăng|Nên giảm|Tăng cường quản lý|Giảm chi phí|Khuyến khích)\s*[:\-]?",
        ]
        is_hallucination = any(re.search(pat, line, re.IGNORECASE) for pat in hallucination_patterns)
        if is_hallucination:
            filtered_lines["hallucination"].append(line[:80])
            continue

        # Remove title-like lines that restate document headers
        if re.search(r"^(Bao cao|Tong ket|De an|Du an)\b", line, re.IGNORECASE) and len(line.split()) <= 12:
            filtered_lines["heading"].append(line)
            continue

        lines.append(line)

    # Collapse into single paragraph with sentence-space separation
    result = " ".join(lines)

    # Clean up excessive spacing
    result = re.sub(r"  +", " ", result).strip()

    # Soft trim to sentence boundary if still too long
    words = result.split()
    if len(words) > 180:
        sentences = re.split(r"(?<=[.!?])\s+", result)
        trimmed: list[str] = []
        word_count = 0
        for sentence in sentences:
            sentence_words = sentence.split()
            if not sentence_words:
                continue
            next_count = word_count + len(sentence_words)
            if next_count > 180:
                break
            trimmed.append(sentence)
            word_count = next_count
            if word_count >= 120:
                # Stop once we have a complete short summary
                break
        if trimmed:
            result = " ".join(trimmed).strip()
            log.debug(f"Summary postprocess: Soft-trimmed to {word_count} words")

    if filtered_lines["leakage"]:
        log.debug(f"Filtered prompt leakage: {filtered_lines['leakage']}")
    if filtered_lines["heading"]:
        log.debug(f"Filtered headings: {filtered_lines['heading']}")
    if filtered_lines["hallucination"]:
        log.debug(f"Filtered hallucinations: {filtered_lines['hallucination']}")
    
    log.debug(f"Summary postprocess: After filtering ({len(result)} chars):\n{result[:500]}")

    # Ensure output is not empty
    if not result:
        log.warning(f"Summary postprocess: All content filtered. Leakage lines: {len(filtered_lines['leakage'])}, Headings: {len(filtered_lines['heading'])}, Hallucinations: {len(filtered_lines['hallucination'])}")
        return "(No substantive content extracted from source document)"

    if _looks_like_refusal(result):
        return result

    return result


def _looks_like_refusal(text: str) -> bool:
    lowered = (text or "").lower()
    return any(
        signal in lowered
        for signal in (
            "xin lỗi",
            "không đủ",
            "không thể",
            "câu trả lời trước",
            "không có đủ thông tin",
            "không đủ để",
        )
    )


def _looks_like_source_copy(summary: str, source: str) -> bool:
    summary_clean = re.sub(r"\s+", " ", (summary or "").lower()).strip()
    source_clean = re.sub(r"\s+", " ", (source or "").lower()).strip()
    if not summary_clean or not source_clean:
        return False
    if len(summary_clean) > 900:
        return True
    if len(summary_clean) > 300 and summary_clean in source_clean:
        return True

    words = re.findall(r"\w+", summary_clean, flags=re.UNICODE)
    if len(words) < 35:
        return False
    spans = [" ".join(words[i : i + 18]) for i in range(0, max(len(words) - 17, 0), 6)]
    copied_spans = sum(1 for span in spans if span and span in source_clean)
    return bool(spans) and copied_spans / len(spans) >= 0.45


def _fallback_summary_from_source(text: str, max_chars: int = 900) -> str:
    cleaned = _prepare_summary_source_text(text)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return "Khong co du noi dung de tom tat."

    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+|\n+", cleaned)
        if len(sentence.strip()) >= 30
    ]
    if not sentences:
        return cleaned[:max_chars].strip()

    keywords = (
        "quyet dinh", "thong bao", "bao cao", "ke hoach", "muc tieu", "yeu cau",
        "nhiem vu", "noi dung", "ket qua", "doi tuong", "pham vi", "thoi han",
        "ngay", "thang", "nam", "uy ban", "so", "bo", "phong", "ban hanh",
        "trien khai", "thuc hien", "quy dinh", "can cu",
    )

    def score(sentence: str, index: int) -> int:
        lowered = sentence.lower()
        keyword_score = sum(3 for keyword in keywords if keyword in lowered)
        number_score = min(len(re.findall(r"\d+", sentence)), 4)
        length_penalty = 4 if len(sentence.split()) > 55 else 0
        position_score = 2 if index < 8 else 0
        return keyword_score + number_score + position_score - length_penalty

    ranked = sorted(enumerate(sentences), key=lambda item: score(item[1], item[0]), reverse=True)
    chosen_indices = sorted(index for index, _ in ranked[:4])
    summary = " ".join(sentences[index] for index in chosen_indices)
    summary = re.sub(r"\s+", " ", summary).strip()
    words = summary.split()
    if len(words) > 130:
        summary = " ".join(words[:130]).rstrip(" ,;:") + "."
    return summary[:max_chars].strip()


def _rebuild_vector_storage(db: Session) -> None:
    rows = db.scalars(
        select(ChunkMetadata)
        .join(Document, ChunkMetadata.document_id == Document.id)
        .where(Document.deleted_at.is_(None))
        .order_by(ChunkMetadata.document_id.asc(), ChunkMetadata.chunk_index.asc())
    ).all()

    if not rows:
        ocr_service.reset_storage()
        return

    chunks = []
    for row in rows:
        metadata = _json_loads(row.metadata_json, default={})
        metadata.setdefault("document_id", row.document_id)
        metadata.setdefault("chunk_index", row.chunk_index)
        metadata.setdefault("page_number", row.page_number)
        metadata.setdefault("start_line", row.start_line)
        metadata.setdefault("end_line", row.end_line)
        metadata.setdefault("section_type", row.section_type)
        chunks.append({"content": row.bm25_text or row.content_preview or "", "metadata": metadata})

    ocr_service.reset_storage()
    ocr_service.store_embeddings(ocr_service.embed_chunks(chunks))
